from __future__ import annotations

import os
import shutil
from io import BytesIO
from pathlib import Path
from typing import BinaryIO

from docx import Document
from pypdf import PdfReader

try:
    from PIL import Image, ImageEnhance, ImageFilter, ImageOps
except ImportError:  # pragma: no cover
    Image = ImageEnhance = ImageFilter = ImageOps = None  # type: ignore[assignment]

try:
    import pytesseract
except ImportError:  # pragma: no cover
    pytesseract = None  # type: ignore[assignment]


IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".tif", ".tiff", ".bmp"}


class DocumentReadError(Exception):
    pass


def _max_upload_bytes() -> int:
    try:
        megabytes = max(1, int(os.getenv("MAX_UPLOAD_MB", "15")))
    except ValueError:
        megabytes = 15
    return megabytes * 1024 * 1024


def _validate_size(data: bytes) -> None:
    limit = _max_upload_bytes()
    if len(data) > limit:
        raise DocumentReadError(
            f"حجم فایل بیشتر از حد مجاز است. حداکثر حجم فعلی {limit // (1024 * 1024)} مگابایت است."
        )


def _validate_extracted_text(text: str) -> str:
    limit = max(1000, int(os.getenv("MAX_EXTRACTED_CHARS", "250000")))
    if len(text) > limit:
        raise DocumentReadError(
            f"متن استخراج شده بیش از حد مجاز است. حداکثر تعداد نویسه فعلی {limit:,} است."
        )
    return text.strip()


def is_image_filename(filename: str) -> bool:
    return Path(filename.lower()).suffix in IMAGE_EXTENSIONS


def _configure_tesseract() -> str | None:
    if pytesseract is None:
        return None
    configured = os.getenv("TESSERACT_CMD", "").strip()
    candidates = [
        configured,
        shutil.which("tesseract") or "",
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            pytesseract.pytesseract.tesseract_cmd = candidate
            return candidate
    return None


def get_ocr_status() -> tuple[bool, str]:
    if Image is None:
        return False, "کتابخانه Pillow نصب نیست."
    if pytesseract is None:
        return False, "کتابخانه pytesseract نصب نیست."
    if _configure_tesseract() is None:
        return False, "برنامه Tesseract پیدا نشد؛ مسیر آن را در TESSERACT_CMD تنظیم کنید."
    languages = os.getenv("OCR_LANGUAGES", "fas+eng")
    return True, f"OCR آماده است؛ زبان‌های فعال: {languages}"


def _prepare_image_for_ocr(data: bytes):
    if Image is None or ImageOps is None or ImageEnhance is None or ImageFilter is None:
        raise DocumentReadError("برای OCR باید Pillow نصب باشد.")
    try:
        image = Image.open(BytesIO(data))
        image = ImageOps.exif_transpose(image).convert("RGB")
        if image.width < 1600:
            scale = 1600 / max(image.width, 1)
            image = image.resize((1600, max(1, int(image.height * scale))), Image.Resampling.LANCZOS)
        grayscale = ImageOps.grayscale(image)
        grayscale = ImageOps.autocontrast(grayscale)
        grayscale = ImageEnhance.Contrast(grayscale).enhance(1.35)
        return grayscale.filter(ImageFilter.SHARPEN)
    except Exception as exc:  # noqa: BLE001
        raise DocumentReadError(f"تصویر قابل پردازش نیست: {exc}") from exc


def _ocr_pil_image(image) -> str:
    if pytesseract is None or _configure_tesseract() is None:
        raise DocumentReadError("Tesseract OCR در سیستم آماده نیست.")
    languages = os.getenv("OCR_LANGUAGES", "fas+eng").strip() or "fas+eng"
    try:
        return pytesseract.image_to_string(image, lang=languages, config="--oem 3 --psm 6").strip()
    except pytesseract.TesseractError as exc:
        message = str(exc)
        if "Failed loading language" in message or "Error opening data file" in message:
            raise DocumentReadError("فایل زبان fas.traineddata یا eng.traineddata در tessdata پیدا نشد.") from exc
        raise DocumentReadError(f"خطا در OCR: {exc}") from exc


def extract_text_from_image(file: BinaryIO | bytes) -> str:
    data = file if isinstance(file, bytes) else file.read()
    _validate_size(data)
    return _ocr_pil_image(_prepare_image_for_ocr(data))


def _ocr_scanned_pdf(data: bytes) -> str:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return ""
    if Image is None:
        return ""
    max_pages = max(1, int(os.getenv("OCR_MAX_PDF_PAGES", "20")))
    texts: list[str] = []
    try:
        document = fitz.open(stream=data, filetype="pdf")
        for page_index in range(min(document.page_count, max_pages)):
            page = document.load_page(page_index)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.open(BytesIO(pixmap.tobytes("png")))
            prepared = _prepare_image_for_ocr(_pil_to_bytes(image))
            text = _ocr_pil_image(prepared)
            if text:
                texts.append(text)
        document.close()
    except Exception as exc:  # noqa: BLE001
        raise DocumentReadError(f"OCR فایل PDF اسکن‌شده ناموفق بود: {exc}") from exc
    return "\n\n".join(texts).strip()


def _pil_to_bytes(image) -> bytes:
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def extract_text_from_pdf(file: BinaryIO | bytes) -> str:
    data = file if isinstance(file, bytes) else file.read()
    _validate_size(data)
    try:
        reader = PdfReader(BytesIO(data))
        pages_text = [(page.extract_text() or "").strip() for page in reader.pages]
        text = "\n\n".join(item for item in pages_text if item).strip()
        minimum_chars = max(1, int(os.getenv("PDF_TEXT_MIN_CHARS", "40")))
        ocr_enabled = os.getenv("OCR_SCANNED_PDF", "true").strip().lower() in {"1", "true", "yes", "on"}
        if len(text) >= minimum_chars or not ocr_enabled:
            return text
        ocr_text = _ocr_scanned_pdf(data)
        return ocr_text or text
    except DocumentReadError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise DocumentReadError(f"خطا در خواندن PDF: {exc}") from exc


def extract_text_from_docx(file: BinaryIO | bytes) -> str:
    try:
        stream = BytesIO(file) if isinstance(file, bytes) else file
        document = Document(stream)
        parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
    except Exception as exc:  # noqa: BLE001
        raise DocumentReadError(f"خطا در خواندن Word/DOCX: {exc}") from exc


def extract_text_from_upload(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    filename = uploaded_file.name.lower()
    suffix = Path(filename).suffix
    data = uploaded_file.getvalue()
    _validate_size(data)
    if suffix == ".pdf":
        text = extract_text_from_pdf(data)
    elif suffix == ".docx":
        text = extract_text_from_docx(data)
    elif suffix == ".txt":
        text = data.decode("utf-8", errors="ignore").strip()
    elif suffix in IMAGE_EXTENSIONS:
        text = extract_text_from_image(data)
    else:
        raise DocumentReadError("فرمت فایل پشتیبانی نمی‌شود. فرمت‌های مجاز: PDF، DOCX، TXT و تصویر.")
    return _validate_extracted_text(text)
