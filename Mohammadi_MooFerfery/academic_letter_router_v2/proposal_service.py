from __future__ import annotations

from io import BytesIO

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from models import AnalysisResult, LetterProposal
from text_utils import normalize_persian_text


class ProposalService:
    def __init__(self, ollama_service=None) -> None:
        self.ollama_service = ollama_service

    @staticmethod
    def _people_evidence(ranked_people: pd.DataFrame | None, limit: int = 3) -> str:
        if ranked_people is None or ranked_people.empty:
            return ""
        lines: list[str] = []
        for _, row in ranked_people.head(limit).iterrows():
            reasons = row.get("explanations", [])
            if not isinstance(reasons, list):
                reasons = [str(reasons)]
            lines.append(
                f"- {row.get('person_name', 'نامشخص')} | امتیاز {float(row.get('routing_score', 0)):.1f} | "
                + "؛ ".join(str(item) for item in reasons[:3])
            )
        return "\n".join(lines)

    @staticmethod
    def resolve_recipient(audience: str, result: AnalysisResult, ranked_people: pd.DataFrame | None) -> str:
        if audience == "person" and ranked_people is not None and not ranked_people.empty:
            return f"استاد محترم {ranked_people.iloc[0]['person_name']}"
        if audience == "field" and result.field_matches:
            selected_name = result.analysis.selected_fields[0] if result.analysis.selected_fields else ""
            top = next(
                (item for item in result.field_matches if item.canonical_name == selected_name),
                result.field_matches[0],
            )
            department = top.departments[0] if top.departments else top.canonical_name
            return f"مدیریت محترم گروه {department}"
        return "مسئول محترم ارجاع مکاتبات دانشگاه"

    def generate(
        self,
        *,
        original_text: str,
        result: AnalysisResult,
        ranked_people: pd.DataFrame | None,
        use_ollama: bool,
        tone: str,
        mode: str,
        audience: str,
    ) -> LetterProposal:
        recipient = self.resolve_recipient(audience, result, ranked_people)
        if use_ollama and self.ollama_service is not None:
            return self.ollama_service.generate_proposal(
                letter_text=original_text,
                analysis=result.analysis,
                field_matches=result.field_matches,
                people_evidence=self._people_evidence(ranked_people),
                tone=tone,
                mode=mode,
                recipient=recipient,
            )
        return self._template_proposal(
            original_text=original_text,
            result=result,
            recipient=recipient,
            tone=tone,
            mode=mode,
        )

    @staticmethod
    def _template_proposal(
        *,
        original_text: str,
        result: AnalysisResult,
        recipient: str,
        tone: str,
        mode: str,
    ) -> LetterProposal:
        analysis = result.analysis
        field_name = analysis.academic_field if analysis.academic_field not in {"", "Unknown", "نامشخص"} else "موضوع مطرح‌شده"
        subject = f"درخواست بررسی و ارجاع در حوزه {field_name}"
        keywords = "، ".join(analysis.keywords[:6]) or "[موضوعات تخصصی تکمیل شود]"

        if mode == "rewrite" and original_text.strip():
            core = normalize_persian_text(original_text)
            letter = (
                "با سلام و احترام\n\n"
                f"{core}\n\n"
                "خواهشمند است در صورت صلاحدید موضوع بررسی و به واحد یا عضو هیئت علمی مرتبط ارجاع شود.\n\n"
                "با احترام\n"
                "[نام و نام خانوادگی تکمیل شود]\n"
                "[اطلاعات تماس تکمیل شود]"
            )
        else:
            letter = (
                "با سلام و احترام\n\n"
                f"اینجانب در خصوص «{analysis.email_summary or field_name}» درخواست بررسی دارم. "
                f"محورهای تخصصی مرتبط شامل {keywords} است.\n\n"
                "خواهشمند است در صورت صلاحدید این درخواست بررسی و به فرد یا گروه آموزشی مرتبط ارجاع شود.\n\n"
                "با احترام\n"
                "[نام و نام خانوادگی تکمیل شود]\n"
                "[سمت یا وابستگی سازمانی تکمیل شود]\n"
                "[اطلاعات تماس تکمیل شود]"
            )

        return LetterProposal(
            suggested_subject=subject,
            recipient_title=recipient,
            suggested_letter=letter,
            tone=tone,
            improvement_notes=[
                "موضوع نامه به صورت مشخص پیشنهاد شد.",
                "مخاطب و درخواست نهایی شفاف‌تر شده‌اند.",
                "اطلاعات ناموجود به جای حدس زدن با جای خالی مشخص شده‌اند.",
            ],
            missing_information=["نام و نام خانوادگی", "اطلاعات تماس"],
            generation_method="template",
        )


def _set_paragraph_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    properties = paragraph._p.get_or_add_pPr()
    if properties.find(qn("w:bidi")) is None:
        properties.append(OxmlElement("w:bidi"))
    for run in paragraph.runs:
        run.font.name = "Tahoma"
        run_properties = run._element.get_or_add_rPr()
        run_fonts = run_properties.get_or_add_rFonts()
        run_fonts.set(qn("w:cs"), "Tahoma")


def proposal_to_docx_bytes(proposal: LetterProposal) -> bytes:
    document = Document()
    title = document.add_heading(proposal.suggested_subject or "پیش‌نویس پیشنهادی نامه", level=1)
    _set_paragraph_rtl(title)
    if proposal.recipient_title:
        recipient = document.add_paragraph(proposal.recipient_title)
        _set_paragraph_rtl(recipient)
    for text in proposal.suggested_letter.split("\n"):
        paragraph = document.add_paragraph(text)
        _set_paragraph_rtl(paragraph)
    if proposal.improvement_notes:
        notes_heading = document.add_heading("نکات بهبود", level=2)
        _set_paragraph_rtl(notes_heading)
        for note in proposal.improvement_notes:
            paragraph = document.add_paragraph(note, style="List Bullet")
            _set_paragraph_rtl(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
